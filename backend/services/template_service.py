from docx import Document
from io import BytesIO


def create_default_template():
    """
    Creates a Word document with Jinja2 template placeholders
    that can be used with python-docx-template.
    
    Returns:
        BytesIO: A BytesIO buffer containing the template document
    """
    # Create a new document
    doc = Document()
    
    # Background section
    doc.add_heading('Background', level=1)
    doc.add_paragraph('{{ background }}')
    
    # Product section
    doc.add_heading('Product', level=1)
    doc.add_paragraph('{{ product }}')
    
    # Usage section
    doc.add_heading('Usage', level=1)
    doc.add_paragraph('{{ usage }}')
    
    # Model Risk Issues section
    doc.add_heading('Model Risk Issues', level=1)
    
    # Create table for Model Risk Issues
    # The table will have a header row and template rows for dynamic content
    risk_table = doc.add_table(rows=3, cols=1)
    
    # First row - loop start
    risk_table.cell(0, 0).text = '{%tr for model_risk_issue in model_risk_issues %}'
    
    # Second row - content template
    risk_table.cell(1, 0).text = ('{{ model_risk_issue.id }} - {{ model_risk_issue.title }} - '
                                   '{{ model_risk_issue.description }} - '
                                   '{{ model_risk_issue.category }} - {{ model_risk_issue.importance }}')
    
    # Third row - loop end
    risk_table.cell(2, 0).text = '{%tr endfor %}'
    
    # Add some spacing
    doc.add_paragraph()
    
    # Model Limitations section
    doc.add_heading('Model Limitations', level=1)
    
    # Create table for Model Limitations
    limitations_table = doc.add_table(rows=3, cols=1)
    
    # First row - loop start
    limitations_table.cell(0, 0).text = '{%tr for model_limitation in model_limitations %}'
    
    # Second row - content template
    limitations_table.cell(1, 0).text = ('{{ model_limitation.id }} - {{ model_limitation.title }} - '
                                          '{{ model_limitation.description }} - '
                                          '{{ model_limitation.category }}')
    
    # Third row - loop end
    limitations_table.cell(2, 0).text = '{%tr endfor %}'
    
    # Save to BytesIO buffer
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io